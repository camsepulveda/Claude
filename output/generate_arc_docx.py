#!/usr/bin/env python3
"""Generate TCT Oncology branded ARC Volume Impact document (.docx) using official brand colors and logo."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# TCT BRAND COLORS (extracted from logo)
TCT_GREEN = RGBColor(0x80, 0xB0, 0x20)       # Primary green - "TCT" letters
TCT_TEAL = RGBColor(0x50, 0x90, 0xB0)         # Teal arc element
TCT_DARK_GREEN = RGBColor(0x5A, 0x80, 0x16)   # Darker green for headers
TCT_LIGHT_TEAL = RGBColor(0x6A, 0xA0, 0xBE)   # Light teal for accents
TCT_DARK = RGBColor(0x2C, 0x2C, 0x2C)         # Body text
TCT_GRAY = RGBColor(0x88, 0x88, 0x88)          # Secondary text
TCT_MED_GRAY = RGBColor(0x66, 0x66, 0x66)      # Medium gray
TCT_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TCT_RED = RGBColor(0xC0, 0x39, 0x2B)
TCT_POS_GREEN = RGBColor(0x27, 0xAE, 0x60)

LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tct_logo.png")

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, color="D0D0D0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>')
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>')
    tcPr.append(tcMar)

def add_styled_table(doc, headers, rows, header_color=None):
    hdr_hex = header_color or "5090B0"
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.space_before = Pt(1)
        p.space_after = Pt(1)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = TCT_WHITE
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, hdr_hex)
        set_cell_border(cell, hdr_hex)
        set_cell_margins(cell)

    for r_idx, row_data in enumerate(rows):
        is_summary = any(k in str(row_data[0]).lower() for k in ["total", "capture", "treated", "access", "growth", "on-island"])
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.space_before = Pt(1)
            p.space_after = Pt(1)
            text = str(val)
            run = p.add_run(text)
            run.font.size = Pt(7.5)
            run.font.name = "Calibri"

            if c_idx == 0:
                run.bold = True
                run.font.color.rgb = TCT_DARK if not is_summary else TCT_TEAL
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if "+" in text and any(c.isdigit() for c in text):
                    run.font.color.rgb = TCT_POS_GREEN
                    run.bold = True
                elif text.startswith("-") and any(c.isdigit() for c in text):
                    run.font.color.rgb = TCT_RED
                elif is_summary:
                    run.bold = True
                    run.font.color.rgb = TCT_TEAL
                else:
                    run.font.color.rgb = TCT_DARK

            bg = "F7F9FA" if r_idx % 2 == 1 else "FFFFFF"
            if is_summary:
                bg = "EAF4F0"
            set_cell_shading(cell, bg)
            set_cell_border(cell)
            set_cell_margins(cell)

    return table

def add_section_header(doc, number, title):
    p = doc.add_paragraph()
    p.space_before = Pt(16)
    p.space_after = Pt(2)
    run = p.add_run(f"{number}  ")
    run.font.size = Pt(9)
    run.font.color.rgb = TCT_GREEN
    run.font.name = "Calibri"
    run.bold = True
    run = p.add_run("·  ")
    run.font.size = Pt(9)
    run.font.color.rgb = TCT_GRAY
    run.font.name = "Calibri"
    run = p.add_run(title.upper())
    run.font.size = Pt(13)
    run.font.color.rgb = TCT_DARK
    run.font.name = "Calibri"
    run.bold = True
    # Green underline
    p2 = doc.add_paragraph()
    p2.space_after = Pt(4)
    pPr = p2._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="80B020"/></w:pBdr>')
    pPr.append(pBdr)

def add_subsection(doc, title):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(10)
    run.font.color.rgb = TCT_TEAL
    run.font.name = "Calibri"
    run.bold = True

def add_body(doc, text, bold=False, color=None, size=9, space_after=3):
    p = doc.add_paragraph()
    p.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.font.color.rgb = color or TCT_DARK
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3)
    p.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"
        run.font.color.rgb = TCT_GREEN
        run = p.add_run(f" — {text}")
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"
        run.font.color.rgb = TCT_DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(8.5)
        run.font.name = "Calibri"
        run.font.color.rgb = TCT_DARK

def add_callout(doc, text, color=None):
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="80B020"/></w:pBdr>')
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(8.5)
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = color or TCT_DARK_GREEN

# ============================================================
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)
style.font.color.rgb = TCT_DARK

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ============================================================
# HEADER WITH LOGO
# ============================================================
if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(2.8))
    p.space_after = Pt(6)

# Teal divider
p = doc.add_paragraph()
p.space_after = Pt(10)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="10" w:space="1" w:color="5090B0"/></w:pBdr>')
pPr.append(pBdr)

# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.space_after = Pt(2)
run = p.add_run("BMTCI PROGRAM")
run.font.size = Pt(22)
run.font.color.rgb = TCT_GREEN
run.font.name = "Calibri"
run.bold = True

p = doc.add_paragraph()
p.space_after = Pt(3)
run = p.add_run("Volume Impact Analysis")
run.font.size = Pt(15)
run.font.color.rgb = TCT_TEAL
run.font.name = "Calibri"

p = doc.add_paragraph()
p.space_after = Pt(3)
run = p.add_run("Prepared for American Red Cross (ARC) — Contract Review & Scope Expansion")
run.font.size = Pt(9.5)
run.font.color.rgb = TCT_DARK
run.font.name = "Calibri"
run.italic = True

p = doc.add_paragraph()
p.space_after = Pt(1)
run = p.add_run("In Collaboration with Centro Comprensivo de Cancer · Universidad de Puerto Rico")
run.font.size = Pt(8.5)
run.font.color.rgb = TCT_GRAY
run.font.name = "Calibri"

p = doc.add_paragraph()
p.space_after = Pt(1)
run = p.add_run("TCT Participation at CCCUPR Effective August 2026  |  Confidential")
run.font.size = Pt(8.5)
run.font.color.rgb = TCT_GRAY
run.font.name = "Calibri"

p = doc.add_paragraph()
p.space_after = Pt(6)
run = p.add_run("June 2026")
run.font.size = Pt(8.5)
run.font.color.rgb = TCT_GRAY
run.font.name = "Calibri"

# ============================================================
# 01 PURPOSE
# ============================================================
add_section_header(doc, "01", "Purpose")
add_body(doc, "This document demonstrates:")
add_bullet(doc, "The projected apheresis collection and cell processing volumes that the BMTCI Program at CCCUPR will generate under the ARC services agreement, justifying contract scope expansion to include autologous HSC collection, CAR-T lymphocyte apheresis, and allogeneic donor collections.", bold_prefix="For ARC")
add_bullet(doc, "How TCT's participation at CCCUPR (effective August 2026) will redistribute hematologic oncology patient volumes across Puerto Rico's hospital landscape, retaining patients on-island who currently travel to the US mainland.", bold_prefix="For Stakeholders")

# ============================================================
# 02 PROGRAM OVERVIEW
# ============================================================
add_section_header(doc, "02", "Program Overview")
add_body(doc, "The BMTCI Program (Hematopoietic Stem Cell Transplantation & Cellular Immunotherapy) launches at CCCUPR as a new department — peer to Medical Oncology and Surgical Oncology — under the Hospital Medical Director. TCT Oncology provides transplant/cellular therapy expertise in collaboration with the CCCUPR academic platform.")

add_subsection(doc, "TCT Oncology Dual-Site Presence")
add_body(doc, "TCT operates at both CCCUPR (BMTCI Program) and Auxilio Mutuo Hospital (HAM), where TCT has an established allogeneic HSCT program and CIBMTR reporting infrastructure. This creates a complementary referral network — CCCUPR handles auto-HCT and CAR-T, while allo-HSCT patients identified at CCCUPR are referred to HAM through TCT, increasing HAM's allogeneic volume.")

add_styled_table(doc,
    ["Phase", "Timeline", "Scope", "Volume Target"],
    [
        ["Phase 1", "Year 1 (2026-2027)", "Autologous HCT", "10-15 auto-HCTs"],
        ["Phase 2", "Year 2-3 (2028-2029)", "+ CAR-T, donor search, HLA", "20-30+ transplants/yr"],
        ["Phase 3", "Year 3+ (2029+)", "+ Allo HCT, gene therapy, FACT", "Full-spectrum cellular therapy"],
    ], header_color="80B020")

add_body(doc, "")
add_body(doc, "Infrastructure: 3-4 HEPA-filtered rooms, outpatient transplant clinic, 2-3 apheresis chairs, pharmacy BSH, ICU with CRS/neurotoxicity capability. Staffing: 18-22 FTE (Phase 1), 26-30 (Phase 2), 32-37 (Phase 3).", size=7.5, color=TCT_GRAY)

# ============================================================
# 03 DISEASE BURDEN
# ============================================================
add_section_header(doc, "03", "Puerto Rico Hematologic Malignancy Burden")
add_subsection(doc, "Annual Disease Incidence")

add_styled_table(doc,
    ["Cancer Type", "Incidence (per 100K)", "Est. Cases/Year", "Notes"],
    [
        ["Leukemia (all subtypes)", "10.2", "~325", "9th most common cancer in PR"],
        ["Non-Hodgkin Lymphoma", "~13.5", "~430", "6th males, 7th females"],
        ["Hodgkin Lymphoma", "~2.5", "~80", "—"],
        ["Multiple Myeloma", "~6.5", "~210", "Increasing, younger women"],
        ["Other (MDS, MPN)", "~2.0", "~65", "—"],
        ["Total", "~34.7", "~1,110", ""],
    ])
add_body(doc, "Sources: PR Central Cancer Registry; SEER Program; Pop: 3,184,835 (US Census 2025)", size=7, color=TCT_GRAY)

add_subsection(doc, "Treatment-Eligible Patient Funnel")
add_styled_table(doc,
    ["Stage", "Patients/Year", "Notes"],
    [
        ["New hematologic malignancies", "~1,110", "All subtypes"],
        ["Auto SCT-eligible (MM, NHL, HD)", "~150-200", "Consolidation or salvage"],
        ["Allo-HSCT eligible (AML, ALL, MDS)", "~40-60", "Requires matched donor"],
        ["CAR-T eligible (2nd/3rd+ line)", "~80-120", "After >=1 prior line"],
        ["Currently receiving auto SCT on-island", "~20-35", "HAM only"],
        ["Currently receiving allo-HSCT on-island", "~8-15", "HAM only (limited)"],
        ["Currently receiving CAR-T on-island", "0", "No program exists"],
        ["Traveling to mainland US", "~60-100", "Those who can afford/access"],
        ["Eligible but untreated (access gap)", "~70-120", "Cannot afford or access"],
    ])

# ============================================================
# 04 ARC VOLUMES
# ============================================================
add_section_header(doc, "04", "ARC Services — Projected Collection Volumes")
add_body(doc, "The ARC contract must cover the following services for the BMTCI Program. These volumes represent new activity generated by the program — not redistribution of existing ARC workload.")

add_styled_table(doc,
    ["ARC Service", "Year 1 (2027)", "Year 2 (2028)", "Year 3 (2029)", "Year 5 (2031)"],
    [
        ["Autologous HSC collections", "12-18", "22-32", "30-42", "40-55"],
        ["CAR-T lymphocyte apheresis", "4-8", "10-14", "16-22", "30-45"],
        ["Allo donor collections (HAM via TCT)", "4-6", "8-12", "10-15", "15-22"],
        ["Cryopreservation events", "16-24", "30-44", "40-57", "55-77"],
        ["Product release testing", "20-32", "40-58", "56-80", "85-122"],
        ["Chain-of-custody/identity", "20-32", "40-58", "56-80", "85-122"],
        ["Total apheresis procedures", "20-32", "40-58", "56-80", "85-122"],
    ], header_color="80B020")

add_body(doc, "")
add_subsection(doc, "ARC Contract Scope Requirements")
add_styled_table(doc,
    ["Service Category", "Status", "Action Required"],
    [
        ["Autologous HSC collection & processing", "Confirm", "Verify explicit coverage"],
        ["HSC cryopreservation & storage", "Confirm", "Verify capacity & pricing"],
        ["CAR-T lymphocyte apheresis & shipment", "Likely gap", "Amend contract for leukapheresis"],
        ["Product release testing", "Confirm", "Verify turnaround times"],
        ["Chain-of-identity / custody", "Confirm", "Must meet FACT-JACIE standards"],
        ["Deviation reporting", "Confirm", "Align with CCCUPR quality program"],
        ["Pricing & pass-through costs", "Review", "Negotiate volume-based pricing"],
        ["Capacity & turnaround guarantees", "Critical", "SLA must guarantee emergent slots"],
        ["Allo donor collections (Phase 3)", "Not yet needed", "Plan amendment by Year 2-3"],
    ])

# ============================================================
# 05 HOSPITAL VOLUME IMPACT
# ============================================================
add_section_header(doc, "05", "Hospital Volume Impact")

add_subsection(doc, "Baseline: Where Patients Go Today")
add_styled_table(doc,
    ["Hospital / Destination", "Consults/Yr", "Auto SCT", "Allo-HSCT", "CAR-T"],
    [
        ["CCCUPR", "2,200-2,800", "0", "0", "0"],
        ["Auxilio Mutuo (HAM) — TCT Oncology", "1,800-2,200", "15-25", "8-15", "0"],
        ["HIMA San Pablo Oncologico", "1,500-1,800", "5-10", "0", "0"],
        ["Other PR hospitals & practices", "2,400-3,400", "0", "0", "0"],
        ["Mainland US", "—", "40-60", "20-30", "30-50"],
        ["Access gap (never treated)", "—", "~50-80", "~12-15", "~50-70"],
        ["Total eligible", "~8,000-10,200", "~110-175", "~40-60", "~80-120"],
    ])

add_callout(doc, "On-island treatment rates — Auto SCT: ~18-20%  |  Allo-HSCT: ~20-25%  |  CAR-T: 0%", TCT_RED)

# Auto SCT
add_body(doc, "")
add_subsection(doc, "Autologous SCT — Projected Cases Per Year")
add_styled_table(doc,
    ["Hospital", "Baseline", "Year 1", "Year 2", "Year 3", "Year 5"],
    [
        ["CCCUPR / BMTCI Program", "0", "10-15", "22-30", "30-40", "40-50"],
        ["Auxilio Mutuo (HAM) — TCT", "15-25", "13-22", "12-20", "10-18", "10-15"],
        ["HIMA San Pablo", "5-10", "5-10", "8-12", "10-15", "12-18"],
        ["Mainland US", "40-60", "30-45", "22-35", "15-25", "8-15"],
        ["Access gap", "50-80", "40-60", "28-42", "18-28", "8-12"],
        ["Treated on-island", "20-35", "28-47", "42-62", "50-73", "62-83"],
        ["On-island capture rate", "~18-20%", "~29-31%", "~46-45%", "~60-58%", "~79-75%"],
    ])

# Allo HSCT
add_body(doc, "")
add_subsection(doc, "Allogeneic HSCT — HAM via TCT Oncology")
add_body(doc, "TCT's dual-site presence creates a direct referral pipeline from CCCUPR to HAM for allo-HSCT. Patients identified at CCCUPR who require allogeneic transplant are referred to HAM through TCT — the only team performing allo-HSCT in Puerto Rico.", size=8.5)

add_styled_table(doc,
    ["Hospital", "Baseline", "Year 1", "Year 2", "Year 3", "Year 5"],
    [
        ["Auxilio Mutuo (HAM) — TCT", "8-15", "12-18", "16-24", "20-30", "25-38"],
        ["CCCUPR / BMTCI Program", "0", "0", "0", "0-2", "3-8"],
        ["Mainland US", "20-30", "16-25", "12-20", "10-16", "6-10"],
        ["Access gap", "12-15", "10-12", "8-10", "5-8", "3-5"],
        ["HAM allo growth vs. baseline", "—", "+50-20%", "+100-60%", "+150-100%", "+213-153%"],
    ], header_color="80B020")

add_callout(doc, "CCCUPR diagnoses ~15-25 allo-eligible patients/year who currently go to the mainland or go untreated. TCT at CCCUPR identifies these patients early and routes them to HAM.")

# CAR-T
add_body(doc, "")
add_subsection(doc, "CAR-T Infusions — Projected Cases Per Year")
add_styled_table(doc,
    ["Hospital", "Baseline", "Year 1", "Year 2", "Year 3", "Year 5"],
    [
        ["CCCUPR / BMTCI Program", "0", "4-8", "10-14", "16-22", "30-45"],
        ["Auxilio Mutuo (HAM)", "0", "0", "0", "0", "0"],
        ["HIMA San Pablo", "0", "0", "0", "0", "0"],
        ["Mainland US", "30-50", "25-42", "22-36", "18-28", "10-16"],
        ["Access gap", "50-70", "42-58", "35-48", "28-38", "15-22"],
        ["Actually treated", "30-50", "29-50", "32-50", "34-50", "40-61"],
        ["Treatment access rate", "~38-42%", "~41-46%", "~48-51%", "~55-57%", "~73%"],
    ])

# Consults
add_body(doc, "")
add_subsection(doc, "Hematology/Oncology Consults Per Year")
add_styled_table(doc,
    ["Hospital", "Baseline", "Year 1", "Year 2", "Year 3", "Year 5"],
    [
        ["CCCUPR (incl. BMTCI)", "2,200-2,800", "2,500-3,200", "2,700-3,400", "2,900-3,600", "3,200-3,900"],
        ["Auxilio Mutuo (HAM) — TCT", "1,800-2,200", "1,850-2,250", "1,900-2,300", "1,950-2,350", "2,000-2,400"],
        ["HIMA San Pablo", "1,500-1,800", "1,500-1,800", "1,500-1,800", "1,500-1,800", "1,500-1,800"],
        ["Other PR hospitals", "2,400-3,400", "2,350-3,350", "2,300-3,300", "2,300-3,300", "2,300-3,300"],
    ])
add_callout(doc, "Both CCCUPR and HAM consult volumes increase. The TCT dual-site model is additive, not competitive.")

# ============================================================
# 06 NET IMPACT
# ============================================================
add_section_header(doc, "06", "Net Impact on Each Hospital")

add_styled_table(doc,
    ["Hospital", "Auto SCT", "Allo-HSCT", "CAR-T", "Consults", "Net"],
    [
        ["CCCUPR", "+10-50/yr", "0 to 0-8", "+4-45/yr", "+300-1,100", "Strong positive"],
        ["HAM — TCT", "-2 to -10", "+4 to +23", "No change", "+50-200", "Net positive"],
        ["HIMA San Pablo", "Minimal", "No change", "No change", "Stable", "Neutral"],
        ["Other PR", "No change", "No change", "No change", "-50 to -100", "Minimal"],
        ["Mainland US", "-10 to -45", "-6 to -20", "-5 to -34", "—", "Significant reduction"],
        ["Access gap", "-10 to -68", "-4 to -10", "-8 to -48", "—", "Major positive"],
    ])

add_body(doc, "")
add_subsection(doc, "Auxilio Mutuo (HAM) — Detailed Impact via TCT Oncology")
add_body(doc, "TCT's dual-site model makes HAM a net beneficiary of the BMTCI Program at CCCUPR:", size=8.5)

add_styled_table(doc,
    ["HAM Metric", "Baseline", "Year 5 (2031)", "Change"],
    [
        ["Auto SCT cases/year", "15-25", "10-15", "-5 to -10 (mild decline)"],
        ["Allo-HSCT cases/year", "8-15", "25-38", "+17 to +23 (significant growth)"],
        ["Total transplants/year", "23-40", "35-53", "+12 to +13 net gain"],
        ["Hem/onc consults/year", "1,800-2,200", "2,000-2,400", "+200 (allo workups)"],
        ["CIBMTR reportable cases", "23-40", "35-53", "+52-33% increase"],
    ], header_color="80B020")

add_callout(doc, "The allo-HSCT growth alone more than compensates for any auto SCT volume shift. HAM's transplant program becomes larger and more complex — strengthening its FACT credentials and academic profile.")

# ============================================================
# 07 KEY TAKEAWAY
# ============================================================
add_section_header(doc, "07", "Key Takeaway")

add_body(doc, "The BMTCI Program at CCCUPR does not cannibalize existing Puerto Rico hospital programs.", bold=True, color=TCT_GREEN, size=11)
add_body(doc, "The overwhelming majority of volume comes from three sources:")
add_bullet(doc, "patients who currently have no on-site transplant option", bold_prefix="CCCUPR's own patients")
add_bullet(doc, "patients kept on-island instead of traveling to the mainland", bold_prefix="Mainland repatriation")
add_bullet(doc, "patients who would never have been treated", bold_prefix="Access gap recovery")

add_body(doc, "")
add_subsection(doc, "Transformation of Puerto Rico's Cellular Therapy Landscape")

add_styled_table(doc,
    ["Metric", "Today", "Year 5 (2031)", "Change"],
    [
        ["On-island auto SCT/year", "20-35", "62-83", "+210-137%"],
        ["On-island allo-HSCT (HAM via TCT)", "8-15", "25-38", "+213-153%"],
        ["On-island CAR-T/year", "0", "30-45", "From zero"],
        ["Total on-island transplants + CAR-T", "28-50", "117-166", "+318-232%"],
        ["On-island capture (auto SCT)", "18-20%", "75-79%", "+57 pts"],
        ["On-island capture (allo-HSCT)", "20-25%", "68-62%", "+43 pts"],
        ["Treatment access (CAR-T)", "38-42%", "73%", "+33 pts"],
        ["Patients traveling to mainland", "60-100/yr", "18-31/yr", "-70%"],
        ["Eligible untreated", "120-150/yr", "23-34/yr", "-78%"],
        ["HAM total transplant volume", "23-40/yr", "35-53/yr", "+52-33%"],
        ["ARC apheresis procedures/year", "0", "85-122", "New volume"],
    ], header_color="5090B0")

add_body(doc, "")
add_subsection(doc, "For ARC")
add_body(doc, "This represents a sustained, growing volume of apheresis collections and cell processing services — starting at 20-32 procedures in Year 1 and scaling to 85-122 procedures by Year 5. This includes auto HSC collections (CCCUPR), CAR-T leukapheresis (CCCUPR), and allo donor collections (HAM via TCT). This is entirely new programmatic volume.")

add_subsection(doc, "For HAM")
add_body(doc, "TCT's dual-site model drives a net increase in HAM's total transplant volume (+52-33% by Year 5). The mild decline in auto SCT (-5 to -10 cases) is more than offset by allo-HSCT growth (+17 to +23 cases) generated by CCCUPR referrals through TCT.")

# ============================================================
# FOOTER
# ============================================================
p = doc.add_paragraph()
p.space_before = Pt(20)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="8" w:space="4" w:color="80B020"/></w:pBdr>')
pPr.append(pBdr)

if os.path.exists(LOGO_PATH):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.6))
    p.space_after = Pt(2)

p = doc.add_paragraph()
p.space_after = Pt(1)
run = p.add_run("Prepared for BMTCI Program stakeholders and American Red Cross contract review")
run.font.size = Pt(7.5)
run.font.color.rgb = TCT_GRAY
run.font.name = "Calibri"

p = doc.add_paragraph()
p.space_after = Pt(1)
run = p.add_run("Carlos Mendez, COO — cmendez@tctoncology.com")
run.font.size = Pt(7.5)
run.font.color.rgb = TCT_MED_GRAY
run.font.name = "Calibri"

p = doc.add_paragraph()
run = p.add_run("June 2026  |  Confidential")
run.font.size = Pt(7.5)
run.font.color.rgb = TCT_GRAY
run.font.name = "Calibri"

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TCT_Oncology_ARC_Volume_Impact_Analysis.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
