#!/usr/bin/env python3
"""Generate OncoCel LLC ARC document using official Brand Identity System."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# ONCOCEL BRAND IDENTITY SYSTEM — COLOR PALETTE (Section 05)
# ============================================================
DEEP_NAVY   = RGBColor(0x0C, 0x23, 0x40)  # #0C2340
BIOTECH_BLU = RGBColor(0x1B, 0x5E, 0x8A)  # #1B5E8A
TEAL        = RGBColor(0x00, 0x7C, 0x8A)  # #007C8A
CYAN        = RGBColor(0x00, 0xA3, 0xB5)  # #00A3B5
CYAN_LT     = RGBColor(0x00, 0xBC, 0xD4)  # #00BCD4
GOLD        = RGBColor(0xC4, 0xA3, 0x5A)  # #C4A35A
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BODY_TEXT   = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY  = RGBColor(0x99, 0x99, 0x99)
MED_GRAY    = RGBColor(0x66, 0x66, 0x66)
GREEN_POS   = RGBColor(0x1E, 0x88, 0x49)
RED_NEG     = RGBColor(0xBB, 0x33, 0x33)

# Typography: Instrument Sans (body), DM Mono (labels) — Calibri as Word fallback
FONT_BODY = "Calibri"

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LETTERHEAD = os.path.join(SCRIPT_DIR, "oncocel_letterhead.png")
ICON = os.path.join(SCRIPT_DIR, "oncocel_icon.png")

def shade(cell, h):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{h}"/>'))

def bdr(cell, c="D8D8D8"):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="2" w:space="0" w:color="{c}"/>'
        f'<w:bottom w:val="single" w:sz="2" w:space="0" w:color="{c}"/>'
        f'<w:left w:val="single" w:sz="2" w:space="0" w:color="{c}"/>'
        f'<w:right w:val="single" w:sz="2" w:space="0" w:color="{c}"/>'
        f'</w:tcBorders>'))

def pad(cell, v=45, h=80):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{v}" w:type="dxa"/><w:bottom w:w="{v}" w:type="dxa"/>'
        f'<w:left w:w="{h}" w:type="dxa"/><w:right w:w="{h}" w:type="dxa"/>'
        f'</w:tcMar>'))

def tbl(doc, hdrs, rows, hc="0C2340"):
    t = doc.add_table(rows=1+len(rows), cols=len(hdrs))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = True
    for i, h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; p.space_before = Pt(2); p.space_after = Pt(2)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE; r.font.name = FONT_BODY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade(c, hc); bdr(c, hc); pad(c)
    for ri, rd in enumerate(rows):
        sumrow = any(k in str(rd[0]).lower() for k in ["total","capture","treated","access","growth","on-island","arc "])
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = ""
            p = c.paragraphs[0]; p.space_before = Pt(1); p.space_after = Pt(1)
            tx = str(v); r = p.add_run(tx); r.font.size = Pt(8); r.font.name = FONT_BODY
            if ci == 0:
                r.bold = True; r.font.color.rgb = TEAL if sumrow else DEEP_NAVY
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if "+" in tx and any(d.isdigit() for d in tx): r.font.color.rgb = GREEN_POS; r.bold = True
                elif tx.startswith("-") and any(d.isdigit() for d in tx): r.font.color.rgb = RED_NEG
                elif sumrow: r.bold = True; r.font.color.rgb = TEAL
                else: r.font.color.rgb = BODY_TEXT
            bg = "F4F7F8" if ri % 2 == 1 else "FFFFFF"
            if sumrow: bg = "E6F2F2"
            shade(c, bg); bdr(c); pad(c)

def section(doc, num, title):
    p = doc.add_paragraph(); p.space_before = Pt(18); p.space_after = Pt(2)
    r = p.add_run(f"{num}  "); r.font.size = Pt(9); r.font.color.rgb = GOLD; r.font.name = FONT_BODY; r.bold = True
    r = p.add_run(title.upper()); r.font.size = Pt(13); r.font.color.rgb = DEEP_NAVY; r.font.name = FONT_BODY; r.bold = True
    # Gold underline matching letterhead accent bar
    p2 = doc.add_paragraph(); p2.space_after = Pt(4)
    p2._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="C4A35A"/></w:pBdr>'))

def body(doc, text, bold=False, color=None, size=9, after=3):
    p = doc.add_paragraph(); p.space_after = Pt(after)
    r = p.add_run(text); r.font.size = Pt(size); r.font.name = FONT_BODY; r.bold = bold; r.font.color.rgb = color or BODY_TEXT

def bullet(doc, text, prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3); p.space_after = Pt(2)
    if prefix:
        r = p.add_run(prefix); r.bold = True; r.font.size = Pt(8.5); r.font.name = FONT_BODY; r.font.color.rgb = TEAL
        r = p.add_run(f" — {text}"); r.font.size = Pt(8.5); r.font.name = FONT_BODY; r.font.color.rgb = BODY_TEXT
    else:
        r = p.add_run(text); r.font.size = Pt(8.5); r.font.name = FONT_BODY; r.font.color.rgb = BODY_TEXT

def callout(doc, text):
    p = doc.add_paragraph(); p.space_before = Pt(4); p.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="C4A35A"/></w:pBdr>'))
    r = p.add_run(text); r.font.size = Pt(8.5); r.font.name = FONT_BODY; r.bold = True; r.font.color.rgb = DEEP_NAVY

# ============================================================
doc = Document()
s = doc.styles['Normal']; s.font.name = FONT_BODY; s.font.size = Pt(9); s.font.color.rgb = BODY_TEXT
for sec in doc.sections:
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5); sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.2)

# ============================================================
# LETTERHEAD HEADER (matches brand guide Section 09)
# ============================================================
if os.path.exists(LETTERHEAD):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.space_after = Pt(12)
    run = p.add_run()
    run.add_picture(LETTERHEAD, width=Inches(6.2))

# ============================================================
# TITLE BLOCK
# ============================================================
p = doc.add_paragraph(); p.space_after = Pt(3); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Volume Growth Notification"); r.font.size = Pt(20); r.font.color.rgb = DEEP_NAVY; r.font.name = FONT_BODY; r.bold = True

p = doc.add_paragraph(); p.space_after = Pt(2)
r = p.add_run("Prepared for American Red Cross (ARC)"); r.font.size = Pt(11); r.font.color.rgb = BIOTECH_BLU; r.font.name = FONT_BODY

p = doc.add_paragraph(); p.space_after = Pt(1)
r = p.add_run("Apheresis & Cell Processing Services — Anticipated Volume Increase"); r.font.size = Pt(9); r.font.color.rgb = MED_GRAY; r.font.name = FONT_BODY; r.italic = True

p = doc.add_paragraph(); p.space_after = Pt(8)
r = p.add_run("June 2026  ·  Confidential"); r.font.size = Pt(8.5); r.font.color.rgb = LIGHT_GRAY; r.font.name = FONT_BODY

# ============================================================
# 01 CONTEXT
# ============================================================
section(doc, "01", "Context")

body(doc, "OncoCel LLC is Puerto Rico's first ambulatory bone marrow transplant and cellular therapy center, focused on outpatient BMTCI services and clinical research. Based in San Juan, OncoCel delivers autologous HCT, CAR-T cellular therapy, and — in later phases — gene therapy in an ambulatory setting, with daily monitoring and 24/7 physician availability.")

body(doc, "OncoCel collaborates with TCT Oncology, which provides inpatient transplant services at two hospital sites: the Centro Comprensivo de Cancer de la Universidad de Puerto Rico (CCCUPR), where TCT is launching a new inpatient BMTCI program effective August 2026, and Auxilio Mutuo Hospital (HAM), where TCT operates the established allogeneic HSCT program and CIBMTR reporting infrastructure.")

body(doc, "Together, OncoCel (ambulatory) and TCT Oncology (inpatient at CCCUPR and HAM) create Puerto Rico's first comprehensive cellular therapy network. Both the ambulatory and inpatient programs require ARC apheresis and cell processing services.", after=4)

body(doc, "ARC is the only FACT-accredited apheresis and cell processing provider on the island. As this network scales, ARC's collection and processing volume will increase significantly across all three sites.", after=6)

callout(doc, "Purpose: Provide ARC advance notice of anticipated volume growth driven by OncoCel's ambulatory program and TCT Oncology's inpatient programs at CCCUPR and HAM, enabling capacity planning, contract scope review, and service continuity.")

# ============================================================
# 02 ONCOCEL AMBULATORY PROGRAM
# ============================================================
section(doc, "02", "OncoCel Ambulatory Program")

body(doc, "OncoCel's ambulatory center focuses exclusively on outpatient autologous HCT and research CAR-T. OncoCel does not perform allogeneic transplant or donor search — those services are provided by TCT Oncology at its inpatient hospital sites.", after=4)

tbl(doc,
    ["OncoCel Program", "Year 1", "Year 2"],
    [
        ["Autologous HCT (Phase I)", "8–10", "15–20"],
        ["Research CAR-T", "4–6", "8–10"],
        ["Allogeneic HSCT", "0", "0"],
        ["Donor search", "0", "0"],
        ["Total OncoCel patients", "12–16", "23–30"],
    ], hc="007C8A")

# ============================================================
# 03 ARC VOLUME INCREASE — ONCOCEL ONLY
# ============================================================
section(doc, "03", "ARC Volume Increase — OncoCel Ambulatory")

body(doc, "The following table shows the incremental volume OncoCel will add to ARC's current baseline. These are net-new procedures above existing ARC workload.", bold=True, size=9, color=TEAL, after=6)

tbl(doc,
    ["ARC Service", "+Year 1 (new)", "+Year 2 (new)"],
    [
        ["Autologous HSC collections", "+8–10", "+15–20"],
        ["Research CAR-T leukapheresis", "+4–6", "+8–10"],
        ["Cryopreservation events", "+8–10", "+15–20"],
        ["Product release testing", "+12–16", "+23–30"],
        ["Chain-of-custody / identity", "+12–16", "+23–30"],
        ["Total new ARC procedures (OncoCel)", "+12–16", "+23–30"],
    ], hc="007C8A")

body(doc, "")
body(doc, "Note: All numbers represent OncoCel ambulatory volume only. Inpatient volumes at CCCUPR and HAM (TCT Oncology) are reported separately by those programs.", size=8, color=MED_GRAY)

# ============================================================
# 04 CONTRACT SCOPE
# ============================================================
section(doc, "04", "Contract Scope — Items to Review")

body(doc, "OncoCel requests ARC confirm coverage for the following services:", after=4)

tbl(doc,
    ["Service", "Status", "Action"],
    [
        ["Autologous HSC collection & processing", "Confirm", "Verify explicit coverage in current agreement"],
        ["HSC cryopreservation & long-term storage", "Confirm", "Verify capacity and pricing at projected volume"],
        ["Research CAR-T leukapheresis & shipment", "Likely gap", "Amend contract — new service category required"],
        ["Capacity guarantees & turnaround SLA", "Critical", "Must guarantee slots for emergent collections"],
        ["Volume-based pricing", "Discuss", "As OncoCel + island-wide volume scales"],
    ], hc="007C8A")

# ============================================================
# 05 ISLAND-WIDE IMPACT
# ============================================================
section(doc, "05", "Island-Wide Volume Outlook")

body(doc, "Beyond OncoCel's ambulatory program, the broader cellular therapy network — including TCT Oncology's inpatient programs at CCCUPR and HAM — will transform Puerto Rico's treatment landscape. The table below shows the expected island-wide volume increase as patients who currently fly to the mainland are retained on-island:", after=6)

tbl(doc,
    ["Metric", "Current (Baseline)", "Projected (Year 2)", "Change"],
    [
        ["Auto SCT on-island / year", "20–35", "50–70", "+30–35"],
        ["Research CAR-T on-island / year", "0", "8–10", "+8–10"],
        ["Allo-HSCT on-island / year (HAM)", "8–15", "16–24", "+8–9"],
        ["Total on-island cellular therapy", "28–50", "74–104", "+46–54"],
        ["Patients flying to mainland / year", "60–100", "30–55", "-30–45"],
        ["Total ARC procedures island-wide", "Current baseline", "+35–50 above baseline", "Net new"],
    ], hc="0C2340")

body(doc, "")
callout(doc, "ARC is the sole on-island FACT-accredited provider supporting this growth. As patients are retained in Puerto Rico instead of flying to the mainland, ARC's role becomes increasingly central to the island's cellular therapy infrastructure.")

# ============================================================
# 06 TIMELINE
# ============================================================
section(doc, "06", "Timeline & Next Steps")

tbl(doc,
    ["When", "What", "ARC Impact"],
    [
        ["August 2026", "OncoCel ambulatory center active; TCT inpatient launches at CCCUPR", "First collections anticipated Q4 2026"],
        ["Q4 2026", "First autologous HSC collections (OncoCel)", "+8–10 collections in Year 1"],
        ["Q1 2027", "First research CAR-T leukapheresis (OncoCel)", "New service type — contract amendment needed"],
        ["2028", "Year 2 ramp: 15–20 auto + 8–10 CAR-T (OncoCel)", "+23–30 OncoCel procedures above baseline"],
        ["Ongoing", "Island-wide network scales (OncoCel + TCT)", "+35–50 total island-wide above baseline"],
    ], hc="007C8A")

body(doc, "")
body(doc, "OncoCel proposes a meeting with ARC leadership to:", bold=True, size=9, color=TEAL, after=2)
bullet(doc, "Review current contract scope against projected service needs")
bullet(doc, "Confirm capacity availability for Year 1 collections (Q4 2026 start)")
bullet(doc, "Discuss CAR-T leukapheresis as a new service category and amend agreement")
bullet(doc, "Establish turnaround and capacity SLA for emergent collection slots")
bullet(doc, "Plan volume-based pricing as the program scales")

# ============================================================
# FOOTER — Gold accent bar + OncoCel branding
# ============================================================
p = doc.add_paragraph(); p.space_before = Pt(24)
p._p.get_or_add_pPr().append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="8" w:space="6" w:color="C4A35A"/></w:pBdr>'))

p = doc.add_paragraph(); p.space_after = Pt(0); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("OncoCel"); r.font.size = Pt(14); r.font.color.rgb = DEEP_NAVY; r.font.name = FONT_BODY; r.bold = True

p = doc.add_paragraph(); p.space_after = Pt(1)
r = p.add_run("CELLULAR THERAPY & TRANSPLANTATION"); r.font.size = Pt(7); r.font.color.rgb = LIGHT_GRAY; r.font.name = FONT_BODY

p = doc.add_paragraph(); p.space_after = Pt(1)
r = p.add_run("Carlos E. Méndez Sepúlveda, JD, MHSA, MPH  ·  Co-Founder & CEO"); r.font.size = Pt(8); r.font.color.rgb = MED_GRAY; r.font.name = FONT_BODY

p = doc.add_paragraph(); p.space_after = Pt(1)
r = p.add_run("Bone Marrow Transplant  ·  Cellular Therapy  ·  Research"); r.font.size = Pt(7.5); r.font.color.rgb = LIGHT_GRAY; r.font.name = FONT_BODY; r.italic = True

p = doc.add_paragraph()
r = p.add_run("San Juan, Puerto Rico  ·  June 2026  ·  Confidential"); r.font.size = Pt(7.5); r.font.color.rgb = LIGHT_GRAY; r.font.name = FONT_BODY

# Save
out = os.path.join(SCRIPT_DIR, "OncoCel_ARC_Volume_Notification.docx")
doc.save(out)
print(f"Saved: {out}")
